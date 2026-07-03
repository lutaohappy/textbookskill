#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
教材生成框架 v4.3 —— 模板驱动的通用教材生成器
读取 YAML 课程数据 + 模板 → 自动生成 docx + TOC md。
"""
import sys, os, re, yaml
sys.path.insert(0, os.path.join(os.path.dirname(__file__), '..'))
from shared.docx_utils import (
    create_doc, save_doc, ah, ap, ap_no_indent, add_table, add_table_left, add_code_block,
    add_bullet, sfont, timestamp, add_figure, FONT_HEITI, FONT_SONG
)
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime

_h2 = lambda d, n, t: ah(d, t, 2, n)
_h3 = lambda d, n, t: ah(d, t, 3, n)
_h4 = lambda d, n, t: ah(d, t, 4, n)
_h1 = lambda d, n, t: ah(d, t, 1, '')
_p = lambda d, t: ap(d, t)
_b = lambda d, t: add_bullet(d, t)

TEMPLATE_DIR = os.path.join(os.path.dirname(__file__), 'templates')

class HeadingNumber:
    def __init__(self):
        self.h2 = 0; self.h3 = 0; self.h4 = 0
    def reset_chapter(self):
        self.h2 = 0; self.h3 = 0; self.h4 = 0
    def next(self, level):
        if level == 2: self.h2 += 1; self.h3 = 0; self.h4 = 0; return str(self.h2)
        elif level == 3: self.h3 += 1; self.h4 = 0; return f'{self.h2}.{self.h3}'
        elif level == 4: self.h4 += 1; return f'{self.h2}.{self.h3}.{self.h4}'
        return ''

def load_template(name):
    path = os.path.join(TEMPLATE_DIR, f'{name}.yaml')
    if not os.path.exists(path):
        available = [f.replace('.yaml', '') for f in os.listdir(TEMPLATE_DIR) if f.endswith('.yaml')]
        raise FileNotFoundError(f'模板 "{name}" 不存在。可用: {available}')
    with open(path, encoding='utf-8') as f:
        return yaml.safe_load(f)

def _words(hours, theory_ratio=0.6):
    return int(hours * 45 * 210 * (theory_ratio + (1 - theory_ratio) * 0.3))

def _split_stops(text):
    return [s.strip() + '。' for s in re.split(r'[。！？]', text) if s.strip()]

def _has_content(data, stype):
    if stype == 'knowledge':
        return bool(data.get('topics'))
    return bool(data.get('body') or data.get('title'))

def render_objectives(doc, hn, data, budget):
    _h2(doc, hn, '学习目标')
    goals = data if isinstance(data, dict) else {}
    sections = [('knowledge', '知识目标'), ('skill', '技能目标'), ('quality', '素养目标')]
    for key, label in sections:
        items = goals.get(key, goals.get(key, []))
        if not items:
            items = [f'掌握{label}核心内容']
        _h3(doc, hn, label)
        for item in items:
            _b(doc, item)

def render_case(doc, hn, data, budget):
    title = data.get('title', '教学案例')
    body = data.get('body', '')
    if not body:
        body = '本章通过实际项目案例引入核心知识点，帮助理解该技术在真实工程中的应用场景。'
    _h2(doc, hn, '案例教学')
    _h3(doc, hn, title)
    _p(doc, body)

def render_knowledge_mixed(doc, hn, data, budget, ch_num=1):
    topics = data.get('topics', {})
    if not topics:
        return
    _h2(doc, hn, '核心技术讲解')
    tb = int(budget / max(len(topics), 1))
    for topic_name, tdata in topics.items():
        _h3(doc, hn, topic_name)
        desc = tdata.get('desc', '')
        know = tdata.get('know', '')
        guide = tdata.get('guide', '')
        paras = []
        if desc:
            paras.append(f'在实际项目中，{desc}')
        if know:
            sentences = _split_stops(know)
            tags_pool = [
                ['系统理解这些参数之间的联动关系是进行方案设计的前提。', '这些指标相互制约需要在实际中进行综合权衡。'],
                ['建立清晰的知识框架有助于快速定位和解决问题。', '理解不同类别的差异是技术方案设计的第一步。'],
                ['这是理解后续内容的基础需要在实践中加深认识。', '掌握这一知识点能帮助你在实际项目中做出更准确判断。'],
            ]
            clues = ['分类', '分', '类型', '种类', '体系', '构成']
            relate_clues = ['=', '：', '—', '→', '取决于', '决定', '关系']
            for i, s in enumerate(sentences):
                s = s.strip()
                if not s:
                    continue
                if any(kw in s for kw in relate_clues):
                    paras.append(f'{s}{tags_pool[0][i%2]}')
                elif any(kw in s for kw in clues):
                    paras.append(f'{s}{tags_pool[1][i%2]}')
                else:
                    paras.append(f'{s}{tags_pool[2][i%2]}')
        if guide:
            guide_sentences = _split_stops(guide)
            guide_clues = ['→', '不要', '避免', '不能', '禁止', '错误', '注意']
            warn_pool = [
                ['这是一个需要牢记的关键经验。', '这类连锁反应在实际项目中需要提前预防。'],
                ['这是最常见的问题之一需要逐步培养规范意识。', '养成正确习惯可以避免绝大多数项目风险。'],
            ]
            for j, s in enumerate(guide_sentences):
                s = s.strip()
                if not s:
                    continue
                if any(kw in s for kw in guide_clues):
                    paras.append(f'{s}{warn_pool[0][j%2]}')
                else:
                    paras.append(f'{s}{warn_pool[1][j%2]}')
        base = re.sub(r'[（(].*?[）)]', '', topic_name).strip()
        if not base:
            base = topic_name.split('（')[0].split('(')[0].strip()
        paras.append(f'综上所述，{topic_name}是{base}领域的核心知识模块。掌握这些内容能够直接指导实际项目中的方案设计和模型调优。')
        for para in paras:
            _p(doc, para)

    examples = data.get('code_examples', [])
    for ex in examples:
        _h4(doc, hn, f'代码示例：{ex["title"]}')
        add_code_block(doc, ex['code'])
    diagrams = data.get('diagrams', [])
    for idx, di in enumerate(diagrams, 1):
        _h4(doc, hn, f'示意图：{di["title"]}')
        add_figure(doc, di['path'], f'图{ch_num}-{idx} {di["title"]}')

def render_knowledge_theory(doc, hn, data, budget):
    topics = data.get('topics', {})
    if not topics:
        return
    _h2(doc, hn, '知识讲解')
    for topic_name, tdata in topics.items():
        _h3(doc, hn, topic_name)
        for key, label in [('concept', '概念'), ('principle', '原理'), ('application', '应用'), ('caution', '注意事项')]:
            val = tdata.get(key, '')
            if val:
                _p(doc, f'{label}：{val}')

def render_design(doc, hn, data, budget):
    title = data.get('title', '') or '方案设计任务'
    body = data.get('body', '')
    _h2(doc, hn, '技术方案设计')
    if not body:
        _h3(doc, hn, title)
        _p(doc, '根据本章所学知识设计完整的技术方案。')
        return
    _h3(doc, hn, title)
    for line in body.strip().split('\n'):
        line = line.strip()
        if not line:
            continue
        if line.startswith('设计任务'):
            _p(doc, f'{line}这是核心实践环节需要综合运用本章知识。')
        else:
            _p(doc, line)

def render_comparison(doc, hn, data, budget):
    _h2(doc, hn, '技术选型对比')
    _h3(doc, hn, data.get('title', '方案对比'))
    headers = data.get('headers', [])
    rows = data.get('rows', [])
    if rows:
        if headers and len(headers) == len(rows[0]):
            add_table_left(doc, headers, rows)
        else:
            add_table_left(doc, rows[0], rows[1:])
    if data.get('recommendation'):
        _p(doc, f'选型建议：{data["recommendation"]}')

def render_experiment(doc, hn, data, budget):
    _h2(doc, hn, '实操指南')
    bg = data.get('background', data.get('objective', ''))
    if bg:
        _p(doc, bg)
    steps = data.get('steps', [])
    for i, step in enumerate(steps, 1):
        _p(doc, f'步骤 {i}：{step}')
    verification = data.get('verification', data.get('analysis', ''))
    if verification:
        _p(doc, verification)

def render_requirement(doc, hn, data, budget):
    _h2(doc, hn, '需求分析与方案设计')
    bg = data.get('background', '')
    if bg:
        _p(doc, bg)
    pain = data.get('pain_points', '')
    if pain:
        _p(doc, f'用户痛点：{pain}')
    comp = data.get('comparison', {})
    if comp:
        _h3(doc, hn, '方案对比')
        add_table_left(doc, comp.get('headers', []), comp.get('rows', []))
        conclusion = comp.get('conclusion', '')
        if conclusion:
            _p(doc, f'选型结论：{conclusion}')

def render_self_test(doc, hn, plan, data, budget):
    _h2(doc, hn, '自测题')
    test = data.get('test', [])
    if not test:
        test = [f'{i+1}. 以下关于{plan["title"]}的描述正确的是？ A.____ B.____ C.____ D.____' for i in range(3)]
    think = data.get('think', [])
    if not think:
        think = [f'{i+1}. 结合本章所学分析{plan["title"]}在实际项目中的应用要点。' for i in range(2)]
    practice = data.get('practice', '') or '完成本章对比分析任务，提交实训报告。'
    _h3(doc, hn, '一、选择题')
    for q in test:
        _p(doc, q)
    _h3(doc, hn, '二、思考题')
    for q in think:
        _p(doc, q)
    _h3(doc, hn, '三、实操题')
    _p(doc, practice)

SECTION_RENDERERS = {
    'objectives': render_objectives,
    'case': render_case,
    'knowledge': render_knowledge_mixed,
    'knowledge_theory': render_knowledge_theory,
    'design': render_design,
    'comparison': render_comparison,
    'experiment': render_experiment,
    'requirement': render_requirement,
    'self_test': render_self_test,
}

def gen_chapter(ch_num, doc, cfg, template):
    plan = cfg['chapters'][ch_num - 1]
    total_budget = plan.get('words', _words(plan['hours'], plan.get('theory_ratio', 0.6)))
    hn = HeadingNumber()
    plan_num = lambda: hn.h2

    _h1(doc, '', f'第{ch_num}章 {plan["title"]}')

    sec_type_map = {s['heading']: s for s in template.get('sections', [])}

    for sec in plan.get('sections', []):
        heading = sec['heading']
        sec_data = sec.get('data', {})
        sec_def = sec_type_map.get(heading, {})
        stype = sec_def.get('type', '')
        budget_ratio = sec_def.get('budget_ratio', 0)
        if not stype:
            continue

        if sec_def.get('optional', False) and not _has_content(sec_data, stype):
            continue

        sec_budget = int(total_budget * budget_ratio)
        renderer = SECTION_RENDERERS.get(stype)
        if renderer:
            if stype == 'self_test':
                renderer(doc, hn, plan, sec_data, sec_budget)
            elif stype == 'knowledge':
                renderer(doc, hn, sec_data, sec_budget, ch_num=ch_num)
            else:
                renderer(doc, hn, sec_data, sec_budget)

def create_doc():
    return Document()

def make_cover(doc, cfg):
    course = cfg['course']
    p = doc.add_paragraph()
    run = p.add_run(course)
    sfont(run, FONT_HEITI, 22, bold=True)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

def make_toc(doc, cfg):
    _h2(doc, HeadingNumber(), '目录')
    for ch in cfg.get('chapters', []):
        _p(doc, f'第{ch["num"]}章 {ch["title"]}（{ch["hours"]}学时）')

def generate_all(yaml_path):
    with open(yaml_path, encoding='utf-8') as f:
        cfg = yaml.safe_load(f)

    template = load_template(cfg['template'])
    print(f'📘 教材生成 v4.3 - {cfg["course"]}（模板：{template["name"]}）')
    total_h = sum(ch['hours'] for ch in cfg['chapters'])
    total_w = sum(ch.get('words', _words(ch['hours'], ch.get('theory_ratio', 0.6))) for ch in cfg['chapters'])
    print(f'   总学时: {total_h} | 总字数预算: {total_w}')

    # Resolve relative diagram paths against base
    base = cfg.get('base', '')
    if not base or not os.path.isdir(base):
        base = os.path.dirname(os.path.dirname(os.path.dirname(os.path.abspath(yaml_path))))
        cfg['base'] = base
    for ch in cfg['chapters']:
        for sec in ch.get('sections', []):
            sec_def = {s['heading']: s for s in template.get('sections', [])}.get(sec['heading'], {})
            stype = sec_def.get('type', '')
            if stype in ('knowledge', 'knowledge_theory'):
                diagrams = sec.get('data', {}).get('diagrams', [])
                for di in diagrams:
                    p = di.get('path', '')
                    if p and not os.path.isabs(p):
                        di['path'] = os.path.join(base, cfg['course'], p)

    doc = create_doc()
    make_cover(doc, cfg)
    make_toc(doc, cfg)
    for ch in cfg['chapters']:
        ch_num = ch['num']
        print(f'  📝 第{ch_num}章 {ch["title"]} ({ch["hours"]}学时)...')
        gen_chapter(ch_num, doc, cfg, template)

    course = cfg['course']
    ts = timestamp()

    out_name = f'{course}_教材（v4.3完整版）_{ts}.docx'
    out_path = os.path.join(base, course, out_name)
    save_doc(doc, out_path)
    print(f'  ✅ {out_name}')

    md_content = generate_toc_md(cfg, template)
    md_name = f'{course}_目录与字数分配_{ts}.md'
    md_path = os.path.join(base, course, md_name)
    with open(md_path, 'w', encoding='utf-8') as f:
        f.write(md_content)
    print(f'  ✅ {md_name}')

    alloc_doc = generate_hour_allocation_docx(cfg)
    alloc_name = f'{course}_课时分配表_{ts}.docx'
    alloc_path = os.path.join(base, course, alloc_name)
    save_doc(alloc_doc, alloc_path)
    print(f'  ✅ {alloc_name}')

    impl_doc = generate_implementation_plan_docx(cfg, template)
    impl_name = f'{course}_实施方案_{ts}.docx'
    impl_path = os.path.join(base, course, impl_name)
    save_doc(impl_doc, impl_path)
    print(f'  ✅ {impl_name}')

def generate_toc_md(cfg, template):
    lines = [f'# 《{cfg["course"]}》目录与字数分配\n\n']
    lines.append(f'总学时：{sum(ch["hours"] for ch in cfg["chapters"])}\n\n')
    for ch in cfg['chapters']:
        w = ch.get('words', _words(ch['hours'], ch.get('theory_ratio', 0.6)))
        lines.append(f'## 第{ch["num"]}章 {ch["title"]}（{ch["hours"]}学时，{w}字）\n')
        lines.append(f'理论比例：{ch.get("theory_ratio", 0.6)}\n\n')
        for sec in ch.get('sections', []):
            lines.append(f'- {sec["heading"]}\n')
    return ''.join(lines)

def generate_hour_allocation_docx(cfg):
    doc = create_doc()
    _h1(doc, '', f'{cfg["course"]} 课时分配表')
    headers = ['章序号', '章名称', '学时', '理论比例']
    rows = [[ch['num'], ch['title'], ch['hours'], ch.get('theory_ratio', 0.6)] for ch in cfg['chapters']]
    rows.append(['合计', '', sum(ch['hours'] for ch in cfg['chapters']), ''])
    add_table_left(doc, headers, rows)
    return doc

def generate_implementation_plan_docx(cfg, template):
    doc = create_doc()
    _h1(doc, '', f'{cfg["course"]} 教学实施方案')
    _p(doc, f'课程名称：{cfg["course"]}')
    for ch in cfg['chapters']:
        _h2(doc, HeadingNumber(), f'第{ch["num"]}章 {ch["title"]}')
        _p(doc, f'学时：{ch["hours"]}')
    return doc
