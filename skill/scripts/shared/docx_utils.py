#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""docx 构建工具集"""
import os, re, subprocess, tempfile
from docx import Document
from docx.shared import Pt, Cm, Emu, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from datetime import datetime

FONT_HEITI = '黑体'
FONT_SONG = '宋体'
FONT_FANGSONG = '仿宋'

def timestamp():
    return datetime.now().strftime('%Y%m%d_%H%M%S')

def create_doc():
    return Document()

def save_doc(doc, path):
    os.makedirs(os.path.dirname(path), exist_ok=True)
    doc.save(path)

def sfont(run, name=FONT_SONG, size=12, bold=False, color=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), name)
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)

def ah(doc, heading_num, text, level, seq=''):
    p = doc.add_paragraph()
    run = p.add_run(text)
    sizes = {1: 18, 2: 15, 3: 13, 4: 12}
    sfont(run, FONT_HEITI, sizes.get(level, 12), bold=True)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    if level == 1:
        p.space_before = Pt(18)
        p.space_after = Pt(12)
    else:
        p.space_before = Pt(12)
        p.space_after = Pt(6)
    return p

def ap(doc, text, indent=True):
    p = doc.add_paragraph()
    run = p.add_run(text)
    sfont(run, FONT_SONG, 12)
    p.paragraph_format.first_line_indent = Cm(0.74) if indent else None
    p.space_after = Pt(3)
    return p

def ap_no_indent(doc, text):
    return ap(doc, text, indent=False)

def add_code_block(doc, code):
    p = doc.add_paragraph()
    run = p.add_run(code)
    sfont(run, 'Courier New', 9)
    p.paragraph_format.left_indent = Cm(1)
    p.space_before = Pt(6)
    p.space_after = Pt(6)
    return p

def add_bullet(doc, text):
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(text)
    sfont(run, FONT_SONG, 12)
    return p

def add_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            sfont(p.runs[0], FONT_HEITI, 10, bold=True) if p.runs else None
    for r_idx, row in enumerate(rows, 1):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx].cells[c_idx]
            cell.text = str(val)
    return table

def add_table_left(doc, headers, rows):
    return add_table(doc, headers, rows)

def add_figure(doc, img_path, caption=''):
    if not os.path.exists(img_path):
        ap(doc, f'[图片未找到: {img_path}]', indent=False)
        return
    # Try SVG conversion if needed
    if img_path.lower().endswith('.svg'):
        png_path = img_path.replace('.svg', '.png')
        if not os.path.exists(png_path):
            try:
                import cairosvg
                cairosvg.svg2png(url=img_path, write_to=png_path, output_width=800)
            except Exception:
                ap(doc, f'[SVG转换失败: {img_path}]', indent=False)
                return
        img_path = png_path
    try:
        p = doc.add_paragraph()
        run = p.add_run()
        run.add_picture(img_path, width=Cm(14))
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if caption:
            cap = doc.add_paragraph()
            cr = cap.add_run(caption)
            sfont(cr, FONT_SONG, 10)
            cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    except Exception as e:
        ap(doc, f'[图片嵌入失败: {e}]', indent=False)
